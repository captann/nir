import os
import re
from docx import Document
from docx.shared import Pt

def parse_c_file(filepath):
    """Парсит C-файл и извлекает структурированную информацию."""
    with open(filepath, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Извлечение импортов
    imports = re.findall(r'#include\s*[<"](.+)[>"]', content)
    
    # Извлечение констант
    constants = re.findall(r'^#define\s+(\w+)\s+(.+)', content, re.MULTILINE)
    
    # Извлечение структур
    structures = re.findall(r'typedef\s+struct\s*{([^}]+)}[^;]*(\w+);', content, re.DOTALL)
    
    # Извлечение функций
    functions = re.findall(r'(\w+\s+\w+\s*\([^)]*\))\s*{', content)
    
    return {
        'imports': imports,
        'constants': constants,
        'structures': structures,
        'functions': functions,
        'raw_content': content
    }

def create_docx_documentation(filename, parsed_info):
    """Создает документацию в формате .docx."""
    doc = Document()
    
    # Заголовок документа
    doc.add_heading(f'Документация: {filename}', 0)
    
    # Импорты
    doc.add_heading('Импорты', level=1)
    for imp in parsed_info['imports']:
        doc.add_paragraph(imp)
    
    # Константы
    doc.add_heading('Константы', level=1)
    for name, value in parsed_info['constants']:
        doc.add_paragraph(f'{name}: {value}')
    
    # Структуры
    doc.add_heading('Структуры', level=1)
    for struct_def, struct_name in parsed_info['structures']:
        doc.add_heading(struct_name, level=2)
        doc.add_paragraph(struct_def.strip())
    
    # Функции
    doc.add_heading('Функции', level=1)
    for func in parsed_info['functions']:
        doc.add_heading(func, level=2)
    
    # Полный текст файла
    doc.add_heading('Полный текст файла', level=1)
    doc.add_paragraph(parsed_info['raw_content'])
    
    # Сохранение документа
    doc_path = os.path.join('d:/keil/prog/docs', f'{filename}_documentation.docx')
    doc.save(doc_path)
    print(f'Создана документация для {filename}')

def generate_documentation_for_core():
    """Генерирует документацию для всех файлов в папке core."""
    core_path = 'd:/keil/prog/src/core'
    
    for filename in os.listdir(core_path):
        if filename.endswith('.c'):
            filepath = os.path.join(core_path, filename)
            parsed_info = parse_c_file(filepath)
            create_docx_documentation(filename, parsed_info)

if __name__ == '__main__':
    generate_documentation_for_core()
